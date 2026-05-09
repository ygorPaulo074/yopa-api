"""
Integration tests for agent endpoints:
  POST   /agent                  — creation, API Key generation
  GET    /agent                  — authenticated agent data
  GET    /agent/context          — current context with version
  GET    /agent/context/history  — version history
  GET    /agent/metrics          — aggregated session metrics
  PUT    /agent/context          — context update with version increment
  PATCH  /agent                  — name update
  DELETE /agent                  — removes agent and associated data
"""
import io
import uuid
from unittest.mock import patch, MagicMock

import httpx
import pytest


AGENT_PAYLOAD = {
    "name": "Support Bot",
    "owner": "acme_corp",
    "context": {
        "tone": "formal",
        "language": "pt",
        "persona": "Assistente de suporte",
    },
}


class TestCreateAgent:
    def test_returns_201_with_agent_id_and_api_key(self, client):
        resp = client.post("/agent", json=AGENT_PAYLOAD)
        assert resp.status_code == 201
        body = resp.json()
        assert "agent_id" in body
        assert "api_key" in body
        assert "created_at" in body

    def test_api_key_contains_agent_id(self, client):
        resp = client.post("/agent", json=AGENT_PAYLOAD)
        body = resp.json()
        assert body["api_key"].startswith(body["agent_id"] + ".")

    def test_missing_name_returns_422(self, client):
        resp = client.post("/agent", json={"owner": "x", "context": {}})
        assert resp.status_code == 422

    def test_missing_owner_returns_422(self, client):
        resp = client.post("/agent", json={"name": "x", "context": {}})
        assert resp.status_code == 422


class TestGetAgent:
    def test_returns_agent_data(self, client, agent):
        agent_id, _, headers = agent
        resp = client.get("/agent", headers=headers)
        assert resp.status_code == 200
        body = resp.json()
        assert body["agent_id"] == agent_id
        assert body["name"] == "Test Agent"
        assert body["owner"] == "test_owner"

    def test_invalid_key_returns_401(self, client):
        resp = client.get("/agent", headers={"Authorization": "Bearer fake.key"})
        assert resp.status_code == 401

    def test_missing_auth_returns_403(self, client):
        resp = client.get("/agent")
        assert resp.status_code in (401, 403)


class TestGetContext:
    def test_returns_context_with_version(self, client, agent):
        _, _, headers = agent
        resp = client.get("/agent/context", headers=headers)
        assert resp.status_code == 200
        body = resp.json()
        assert body["version"] == 1
        assert body["tone"] == "formal"
        assert body["language"] == "pt"

    def test_context_history_has_one_entry_after_create(self, client, agent):
        _, _, headers = agent
        resp = client.get("/agent/context/history", headers=headers)
        assert resp.status_code == 200
        body = resp.json()
        assert len(body["versions"]) == 1
        assert body["versions"][0]["version"] == 1


class TestUpdateContext:
    def test_increments_version(self, client, agent):
        _, _, headers = agent
        resp = client.put("/agent/context", headers=headers, json={
            "tone": "informal",
            "language": "en",
        })
        assert resp.status_code == 200
        assert resp.json()["version"] == 2

    def test_history_grows_after_update(self, client, agent):
        _, _, headers = agent
        client.put("/agent/context", headers=headers, json={"tone": "informal"})
        resp = client.get("/agent/context/history", headers=headers)
        body = resp.json()
        assert len(body["versions"]) == 2

    def test_changes_field_reflects_updated_keys(self, client, agent):
        _, _, headers = agent
        client.put("/agent/context", headers=headers, json={"tone": "informal"})
        resp = client.get("/agent/context/history", headers=headers)
        versions = resp.json()["versions"]
        latest = next(v for v in versions if v["version"] == 2)
        assert "tone" in latest["changes"]

    def test_context_reflects_new_values(self, client, agent):
        _, _, headers = agent
        client.put("/agent/context", headers=headers, json={"tone": "informal"})
        resp = client.get("/agent/context", headers=headers)
        assert resp.json()["tone"] == "informal"


class TestGetMetrics:
    def test_returns_zero_metrics_with_no_sessions(self, client, agent):
        agent_id, _, headers = agent
        resp = client.get("/agent/metrics", headers=headers)
        assert resp.status_code == 200
        body = resp.json()
        assert body["agent_id"] == agent_id
        assert body["total_sessions"] == 0
        assert body["total_messages"] == 0
        assert body["resolution_rate"] == 0.0
        assert body["escalation_rate"] == 0.0


class TestDeleteAgent:
    def test_returns_deleted_at(self, client, agent):
        _, _, headers = agent
        resp = client.delete("/agent", headers=headers)
        assert resp.status_code == 200
        assert "deleted_at" in resp.json()

    def test_agent_not_accessible_after_delete(self, client, agent):
        _, _, headers = agent
        client.delete("/agent", headers=headers)
        resp = client.get("/agent", headers=headers)
        assert resp.status_code == 401


class TestWebhookDatasource:
    def test_create_agent_with_webhook_datasource(self, client):
        resp = client.post("/agent", json={
            "name": "Webhook Bot",
            "owner": "acme",
            "context": {
                "webhook_datasource": {
                    "url": "https://hooks.example.com/query",
                    "token": "secret123",
                },
            },
        })
        assert resp.status_code == 201

    def test_context_stores_webhook_datasource(self, client):
        resp = client.post("/agent", json={
            "name": "Webhook Bot",
            "owner": "acme",
            "context": {
                "webhook_datasource": {
                    "url": "https://hooks.example.com/query",
                },
            },
        })
        body = resp.json()
        headers = {"Authorization": f"Bearer {body['api_key']}"}
        ctx = client.get("/agent/context", headers=headers).json()
        assert ctx["webhook_datasource"] is not None
        assert "hooks.example.com" in ctx["webhook_datasource"]["url"]

    def test_update_context_with_webhook_datasource(self, client, agent):
        _, _, headers = agent
        resp = client.put("/agent/context", headers=headers, json={
            "webhook_datasource": {
                "url": "https://hooks.example.com/v2",
                "token": "newtoken",
            },
        })
        assert resp.status_code == 200
        ctx = client.get("/agent/context", headers=headers).json()
        assert "v2" in ctx["webhook_datasource"]["url"]


class TestWebhookTool:
    def test_execute_returns_formatted_results(self):
        from src.infrastructure.tools.webhook_tool import WebhookTool

        fake_response = MagicMock()
        fake_response.json.return_value = [
            {"product": "Widget A", "price": "10.00"},
            {"product": "Widget B", "price": "20.00"},
        ]
        fake_response.raise_for_status = MagicMock()

        with patch("httpx.post", return_value=fake_response) as mock_post:
            tool = WebhookTool(url="https://hooks.example.com/query", token="tok")
            result = tool.execute("widget")

        mock_post.assert_called_once()
        call_kwargs = mock_post.call_args
        assert call_kwargs.kwargs["json"] == {"query": "widget"}
        assert call_kwargs.kwargs["headers"]["Authorization"] == "Bearer tok"
        assert "Widget A" in result

    def test_execute_raises_timeout_error(self):
        from src.infrastructure.tools.webhook_tool import WebhookTool

        with patch("httpx.post", side_effect=httpx.TimeoutException("timeout")):
            tool = WebhookTool(url="https://hooks.example.com/query")
            with pytest.raises(TimeoutError):
                tool.execute("test")

    def test_execute_raises_runtime_error_on_http_error(self):
        from src.infrastructure.tools.webhook_tool import WebhookTool

        with patch("httpx.post", side_effect=httpx.HTTPStatusError(
            "500", request=MagicMock(), response=MagicMock()
        )):
            tool = WebhookTool(url="https://hooks.example.com/query")
            with pytest.raises(RuntimeError):
                tool.execute("test")

    def test_execute_no_token_omits_auth_header(self):
        from src.infrastructure.tools.webhook_tool import WebhookTool

        fake_response = MagicMock()
        fake_response.json.return_value = []
        fake_response.raise_for_status = MagicMock()

        with patch("httpx.post", return_value=fake_response) as mock_post:
            tool = WebhookTool(url="https://hooks.example.com/query")
            tool.execute("test")

        headers = mock_post.call_args.kwargs["headers"]
        assert "Authorization" not in headers

    def test_build_tools_includes_webhook_when_configured(self, client, agent, mock_ai):
        from src.infrastructure.tools.webhook_tool import WebhookTool

        _, _, headers = agent
        client.put("/agent/context", headers=headers, json={
            "webhook_datasource": {
                "url": "https://hooks.example.com/query",
            },
        })

        fake_response = MagicMock()
        fake_response.json.return_value = [{"answer": "42"}]
        fake_response.raise_for_status = MagicMock()

        with patch("httpx.post", return_value=fake_response):
            resp = client.post("/chat", headers=headers, json={
                "session_id": str(uuid.uuid4()),
                "message": "What is the answer?",
            })
        assert resp.status_code == 200


class TestParseContext:
    def test_returns_structured_context(self, client, agent, mock_ai):
        import json
        from src.infrastructure.ai.client import AIClient, AIResponse, AIUsage

        _, _, headers = agent
        parsed_json = json.dumps({
            "tone": "formal",
            "language": "pt",
            "persona": "Assistente de suporte técnico",
            "restrictions": {"topics": ["concorrentes"]},
        })
        fake = AIResponse(content=parsed_json, usage=AIUsage(input_tokens=10, output_tokens=20, total_tokens=30))
        with patch.object(AIClient, "complete", return_value=fake):
            resp = client.post(
                "/agent/parse-context",
                headers=headers,
                json={"text": "Sou formal, falo português, não falo de concorrentes."},
            )

        assert resp.status_code == 200
        body = resp.json()
        assert body["context"]["tone"] == "formal"
        assert body["context"]["language"] == "pt"
        assert "concorrentes" in body["context"]["restrictions"]["topics"]

    def test_returns_422_on_invalid_ai_response(self, client, agent):
        from src.infrastructure.ai.client import AIClient, AIResponse, AIUsage

        _, _, headers = agent
        fake = AIResponse(content="not json at all", usage=AIUsage(input_tokens=1, output_tokens=1, total_tokens=2))
        with patch.object(AIClient, "complete", return_value=fake):
            resp = client.post("/agent/parse-context", headers=headers, json={"text": "qualquer coisa"})
        assert resp.status_code == 422

    def test_empty_text_still_processes(self, client, agent):
        from src.infrastructure.ai.client import AIClient, AIResponse, AIUsage

        _, _, headers = agent
        fake = AIResponse(content="{}", usage=AIUsage(input_tokens=1, output_tokens=1, total_tokens=2))
        with patch.object(AIClient, "complete", return_value=fake):
            resp = client.post("/agent/parse-context", headers=headers, json={"text": ""})
        assert resp.status_code == 200
        assert resp.json()["context"] is not None

    def test_requires_auth(self, client):
        resp = client.post("/agent/parse-context", json={"text": "qualquer coisa"})
        assert resp.status_code in (401, 403)


class TestSqlTool:
    def _make_sqlite_tool(self, agent_id: str = "test-agent") -> "SqlTool":
        from src.infrastructure.tools.sql_tool import SqlTool
        from src.infrastructure.security import encrypt_secret
        conn = encrypt_secret("sqlite:///:memory:")
        return SqlTool(connection_string_enc=conn, agent_id=agent_id)

    def test_validate_connection_string_accepts_postgresql(self):
        from src.infrastructure.tools.sql_tool import validate_connection_string
        assert validate_connection_string("postgresql://user:pass@localhost/db") == "postgresql://user:pass@localhost/db"

    def test_validate_connection_string_rejects_unknown_dialect(self):
        from src.infrastructure.tools.sql_tool import validate_connection_string
        with pytest.raises(ValueError, match="não permitido"):
            validate_connection_string("mssql://user:pass@host/db")

    def test_validate_connection_string_rejects_no_scheme(self):
        from src.infrastructure.tools.sql_tool import validate_connection_string
        with pytest.raises(ValueError):
            validate_connection_string("not-a-url")

    def test_select_only_rejects_insert(self):
        from src.infrastructure.tools.sql_tool import _validate_sql
        with pytest.raises(ValueError, match="SELECT"):
            _validate_sql("INSERT INTO t VALUES (1)")

    def test_select_only_rejects_drop(self):
        from src.infrastructure.tools.sql_tool import _validate_sql
        with pytest.raises(ValueError):
            _validate_sql("DROP TABLE users")

    def test_select_only_rejects_multiple_statements(self):
        from src.infrastructure.tools.sql_tool import _validate_sql
        with pytest.raises(ValueError, match="Múltiplos"):
            _validate_sql("SELECT 1; DROP TABLE users")

    def test_select_passes_validation(self):
        from src.infrastructure.tools.sql_tool import _validate_sql
        assert _validate_sql("SELECT id, name FROM users") == "SELECT id, name FROM users"

    def test_execute_returns_results_from_sqlite(self, tmp_path):
        import os
        os.environ["DATA_PATH"] = str(tmp_path)

        from src.infrastructure.tools.sql_tool import SqlTool
        from src.infrastructure.security import encrypt_secret
        from sqlalchemy import create_engine, text as stext

        db_path = tmp_path / "test.db"
        engine = create_engine(f"sqlite:///{db_path}")
        with engine.begin() as conn:
            conn.execute(stext("CREATE TABLE produtos (id INTEGER, nome TEXT)"))
            conn.execute(stext("INSERT INTO produtos VALUES (1, 'Widget A')"))
        engine.dispose()

        tool = SqlTool(
            connection_string_enc=encrypt_secret(f"sqlite:///{db_path}"),
            agent_id="agent-x",
        )
        result = tool.execute("SELECT id, nome FROM produtos")
        assert "Widget A" in result
        assert "id" in result

    def test_execute_raises_value_error_on_invalid_sql(self, tmp_path):
        import os
        os.environ["DATA_PATH"] = str(tmp_path)

        from src.infrastructure.tools.sql_tool import SqlTool
        from src.infrastructure.security import encrypt_secret
        tool = SqlTool(
            connection_string_enc=encrypt_secret("sqlite:///:memory:"),
            agent_id="agent-x",
        )
        with pytest.raises(ValueError):
            tool.execute("DELETE FROM users")

    def test_audit_log_written_on_execute(self, tmp_path):
        import os
        os.environ["DATA_PATH"] = str(tmp_path)

        from src.infrastructure.tools.sql_tool import SqlTool
        from src.infrastructure.security import encrypt_secret
        from sqlalchemy import create_engine, text as stext

        db_path = tmp_path / "audit.db"
        engine = create_engine(f"sqlite:///{db_path}")
        with engine.begin() as conn:
            conn.execute(stext("CREATE TABLE t (x INTEGER)"))
        engine.dispose()

        tool = SqlTool(
            connection_string_enc=encrypt_secret(f"sqlite:///{db_path}"),
            agent_id="audit-agent",
        )
        tool.execute("SELECT x FROM t")

        audit_file = tmp_path / "agents" / "audit-agent" / "sql_audit.jsonl"
        assert audit_file.exists()
        import json
        entry = json.loads(audit_file.read_text().strip())
        assert entry["success"] is True
        assert "SELECT" in entry["sql"]


class TestSqlDatasource:
    def test_create_agent_with_sql_datasource(self, client):
        resp = client.post("/agent", json={
            "name": "SQL Bot",
            "owner": "acme",
            "context": {
                "sql_datasource": {
                    "connection_string": "sqlite:///:memory:",
                    "allowed_tables": [],
                    "max_rows": 20,
                },
            },
        })
        assert resp.status_code == 201

    def test_context_stores_sql_datasource_encrypted(self, client):
        resp = client.post("/agent", json={
            "name": "SQL Bot",
            "owner": "acme",
            "context": {
                "sql_datasource": {
                    "connection_string": "sqlite:///:memory:",
                    "allowed_tables": [],
                    "max_rows": 10,
                },
            },
        })
        api_key = resp.json()["api_key"]
        headers = {"Authorization": f"Bearer {api_key}"}
        ctx = client.get("/agent/context", headers=headers).json()
        assert ctx["sql_datasource"]["connection_string"].startswith("enc:")

    def test_validate_sql_rejects_unsupported_dialect(self, client, agent):
        _, _, headers = agent
        resp = client.post("/agent/validate-sql", headers=headers, json={"connection_string": "mssql://u:p@host/db"})
        assert resp.status_code == 200
        assert resp.json()["valid"] is False
        assert "não permitido" in resp.json()["error"]

    def test_validate_sql_accepts_sqlite_memory(self, client, agent):
        _, _, headers = agent
        resp = client.post("/agent/validate-sql", headers=headers, json={"connection_string": "sqlite:///:memory:"})
        assert resp.status_code == 200
        assert resp.json()["valid"] is True
        assert resp.json()["dialect"] == "sqlite"

    def test_validate_sql_requires_auth(self, client):
        resp = client.post("/agent/validate-sql", json={"connection_string": "sqlite:///:memory:"})
        assert resp.status_code in (401, 403)

    def test_encrypt_secret_roundtrip(self):
        from src.infrastructure.security import encrypt_secret, decrypt_secret
        original = "postgresql://user:pass@localhost/db"
        enc = encrypt_secret(original)
        assert enc.startswith("enc:")
        assert decrypt_secret(enc) == original

    def test_encrypt_secret_idempotent(self):
        from src.infrastructure.security import encrypt_secret
        original = "sqlite:///:memory:"
        enc1 = encrypt_secret(original)
        enc2 = encrypt_secret(enc1)
        assert enc1 == enc2

    def test_mask_connection_string(self):
        from src.infrastructure.security import mask_connection_string
        conn = "postgresql://admin:secret123@db.example.com/mydb"
        masked = mask_connection_string(conn)
        assert "secret123" not in masked
        assert "***" in masked
        assert "admin" in masked


class TestFileExtractor:
    def test_txt_splits_by_double_newline(self):
        from src.infrastructure.ingestion.file_extractor import extract
        content = b"Hello world.\n\nSecond paragraph."
        records = extract(content, "doc.txt")
        assert len(records) == 2
        assert records[0]["text"] == "Hello world."
        assert records[1]["text"] == "Second paragraph."

    def test_txt_ignores_blank_paragraphs(self):
        from src.infrastructure.ingestion.file_extractor import extract
        content = b"First.\n\n\n\nSecond."
        records = extract(content, "doc.txt")
        assert len(records) == 2

    def test_txt_single_paragraph(self):
        from src.infrastructure.ingestion.file_extractor import extract
        content = b"Just one paragraph."
        records = extract(content, "notes.txt")
        assert records == [{"text": "Just one paragraph."}]

    def test_docx_extracts_paragraphs(self):
        from docx import Document
        from src.infrastructure.ingestion.file_extractor import extract
        doc = Document()
        doc.add_paragraph("First paragraph.")
        doc.add_paragraph("Second paragraph.")
        buf = io.BytesIO()
        doc.save(buf)
        records = extract(buf.getvalue(), "report.docx")
        texts = [r["text"] for r in records]
        assert "First paragraph." in texts
        assert "Second paragraph." in texts

    def test_docx_ignores_empty_paragraphs(self):
        from docx import Document
        from src.infrastructure.ingestion.file_extractor import extract
        doc = Document()
        doc.add_paragraph("Content.")
        doc.add_paragraph("")
        buf = io.BytesIO()
        doc.save(buf)
        records = extract(buf.getvalue(), "file.docx")
        assert all(r["text"] for r in records)

    def test_unsupported_extension_raises(self):
        from src.infrastructure.ingestion.file_extractor import extract
        with pytest.raises(ValueError, match="Unsupported file type"):
            extract(b"data", "file.pptx")


class TestKnowledgeBase:
    def _upload(self, client, headers, filename, content, content_type="text/csv"):
        return client.post(
            "/agent/knowledge/upload",
            headers=headers,
            files={"file": (filename, content, content_type)},
        )

    def test_upload_csv_returns_201_with_metadata(self, client, agent):
        _, _, headers = agent
        csv = b"titulo,conteudo\nHorario,Seg a Sex 9h-18h\n"
        resp = self._upload(client, headers, "base.csv", csv)
        assert resp.status_code == 201
        body = resp.json()
        assert body["filename"] == "base.csv"
        assert body["file_type"] == "csv"
        assert body["record_count"] == 1
        assert "file_id" in body
        assert "uploaded_at" in body

    def test_upload_json_array_returns_correct_count(self, client, agent):
        import json
        _, _, headers = agent
        data = json.dumps([{"titulo": "A", "conteudo": "B"}, {"titulo": "C", "conteudo": "D"}]).encode()
        resp = self._upload(client, headers, "records.json", data, "application/json")
        assert resp.status_code == 201
        assert resp.json()["record_count"] == 2

    def test_upload_txt_returns_201(self, client, agent):
        _, _, headers = agent
        resp = self._upload(client, headers, "faq.txt", b"Pergunta um.\n\nPergunta dois.", "text/plain")
        assert resp.status_code == 201
        assert resp.json()["file_type"] == "txt"
        assert resp.json()["record_count"] == 2

    def test_upload_docx_returns_201(self, client, agent):
        from docx import Document
        _, _, headers = agent
        doc = Document()
        doc.add_paragraph("Conteudo do documento.")
        buf = io.BytesIO()
        doc.save(buf)
        resp = self._upload(
            client, headers, "manual.docx", buf.getvalue(),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        assert resp.status_code == 201
        assert resp.json()["file_type"] == "docx"

    def test_upload_unsupported_format_returns_422(self, client, agent):
        _, _, headers = agent
        resp = self._upload(client, headers, "slides.pptx", b"data", "application/octet-stream")
        assert resp.status_code == 422

    def test_upload_requires_auth(self, client):
        csv = b"titulo,conteudo\nA,B\n"
        resp = client.post(
            "/agent/knowledge/upload",
            files={"file": ("base.csv", csv, "text/csv")},
        )
        assert resp.status_code in (401, 403)

    def test_list_empty_initially(self, client, agent):
        _, _, headers = agent
        resp = client.get("/agent/knowledge", headers=headers)
        assert resp.status_code == 200
        assert resp.json()["files"] == []

    def test_list_reflects_uploaded_files(self, client, agent):
        _, _, headers = agent
        self._upload(client, headers, "a.csv", b"titulo,conteudo\nA,B\n")
        self._upload(client, headers, "b.csv", b"titulo,conteudo\nC,D\nE,F\n")
        resp = client.get("/agent/knowledge", headers=headers)
        assert resp.status_code == 200
        files = resp.json()["files"]
        assert len(files) == 2
        names = {f["filename"] for f in files}
        assert names == {"a.csv", "b.csv"}

    def test_list_shows_correct_record_count(self, client, agent):
        _, _, headers = agent
        self._upload(client, headers, "data.csv", b"titulo,conteudo\nA,B\nC,D\nE,F\n")
        resp = client.get("/agent/knowledge", headers=headers)
        assert resp.json()["files"][0]["record_count"] == 3

    def test_list_isolated_between_agents(self, client):
        resp1 = client.post("/agent", json={"name": "Agent1", "owner": "o1", "context": {}})
        resp2 = client.post("/agent", json={"name": "Agent2", "owner": "o2", "context": {}})
        h1 = {"Authorization": f"Bearer {resp1.json()['api_key']}"}
        h2 = {"Authorization": f"Bearer {resp2.json()['api_key']}"}
        self._upload(client, h1, "data.csv", b"titulo,conteudo\nA,B\n")
        resp = client.get("/agent/knowledge", headers=h2)
        assert resp.json()["files"] == []

    def test_delete_removes_file(self, client, agent):
        _, _, headers = agent
        upload = self._upload(client, headers, "del.csv", b"titulo,conteudo\nA,B\n")
        file_id = upload.json()["file_id"]
        resp = client.delete(f"/agent/knowledge/{file_id}", headers=headers)
        assert resp.status_code == 200
        assert resp.json()["deleted"] is True
        files = client.get("/agent/knowledge", headers=headers).json()["files"]
        assert all(f["file_id"] != file_id for f in files)

    def test_delete_nonexistent_returns_404(self, client, agent):
        _, _, headers = agent
        resp = client.delete("/agent/knowledge/nonexistent-id", headers=headers)
        assert resp.status_code == 404

    def test_delete_requires_auth(self, client):
        resp = client.delete("/agent/knowledge/some-id")
        assert resp.status_code in (401, 403)


class TestSoftDeletes:
    def test_deleted_agent_returns_404_on_get(self, client, agent):
        _, _, headers = agent
        client.delete("/agent", headers=headers)
        resp = client.get("/agent", headers=headers)
        assert resp.status_code in (401, 403, 404)

    def test_deleted_agent_returns_401_on_auth(self, client, agent):
        _, _, headers = agent
        client.delete("/agent", headers=headers)
        resp = client.post("/chat", headers=headers, json={
            "session_id": "s1", "user_id": "u1", "message": "hi",
        })
        assert resp.status_code in (401, 403)

    def test_delete_agent_response_has_deleted_at(self, client, agent):
        _, _, headers = agent
        resp = client.delete("/agent", headers=headers)
        assert resp.status_code == 200
        assert "deleted_at" in resp.json()

    def test_deleted_session_returns_404(self, client, agent, mock_ai):
        _, _, headers = agent
        sid = str(uuid.uuid4())
        client.post("/chat", headers=headers, json={
            "session_id": sid, "user_id": "u1", "message": "hi",
        })
        client.post(f"/chat/{sid}/end", headers=headers)
        resp = client.delete(f"/data/chat/{sid}", headers=headers)
        assert resp.status_code == 204
        resp2 = client.get(f"/data/chat/{sid}", headers=headers)
        assert resp2.status_code == 404

    def test_deleted_session_absent_from_list(self, client, agent, mock_ai):
        _, _, headers = agent
        sid = str(uuid.uuid4())
        client.post("/chat", headers=headers, json={
            "session_id": sid, "user_id": "u1", "message": "hi",
        })
        client.post(f"/chat/{sid}/end", headers=headers)
        client.delete(f"/data/chat/{sid}", headers=headers)
        resp = client.get("/data/chat", headers=headers)
        assert resp.status_code == 200
        session_ids = [c["session_id"] for c in resp.json()["chats"]]
        assert sid not in session_ids

    def test_purge_deleted_removes_agent_after_cutoff(self, client, agent, patch_env):
        from datetime import datetime, timezone, timedelta
        from src.infrastructure.persistence.factory import get_driver
        agent_id, _, headers = agent

        client.delete("/agent", headers=headers)

        future = (datetime.now(timezone.utc) + timedelta(days=1)).isoformat()
        result = get_driver().purge_deleted(before=future)
        assert result["agents_purged"] >= 1

    def test_purge_before_cutoff_preserves_agent(self, client, agent, patch_env):
        from datetime import datetime, timezone, timedelta
        from src.infrastructure.persistence.factory import get_driver
        agent_id, _, headers = agent

        client.delete("/agent", headers=headers)

        past = (datetime.now(timezone.utc) - timedelta(days=1)).isoformat()
        result = get_driver().purge_deleted(before=past)
        assert result["agents_purged"] == 0


class TestUrlFetcher:
    def test_fetch_url_html_returns_201(self, client, agent):
        _, _, headers = agent
        html = "<html><body><p>Parágrafo um.</p><p>Parágrafo dois.</p></body></html>"
        with patch("src.infrastructure.ingestion.url_fetcher.httpx.Client") as mock_client_cls:
            mock_resp = MagicMock()
            mock_resp.headers = {"content-type": "text/html"}
            mock_resp.text = html
            mock_resp.raise_for_status = lambda: None
            mock_client_cls.return_value.__enter__ = lambda s: mock_client_cls.return_value
            mock_client_cls.return_value.__exit__ = MagicMock(return_value=False)
            mock_client_cls.return_value.get = MagicMock(return_value=mock_resp)

            resp = client.post("/agent/knowledge/fetch-url", headers=headers, json={"url": "https://example.com"})

        assert resp.status_code == 201
        body = resp.json()
        assert body["file_type"] == "url"
        assert body["record_count"] == 2
        assert "file_id" in body

    def test_fetch_url_rss_returns_records(self, client, agent):
        _, _, headers = agent
        rss = """<?xml version="1.0"?>
        <rss version="2.0"><channel>
          <item><title>Título A</title><description>Resumo A</description></item>
          <item><title>Título B</title><description>Resumo B</description></item>
        </channel></rss>"""
        with patch("src.infrastructure.ingestion.url_fetcher.httpx.Client") as mock_client_cls:
            mock_resp = MagicMock()
            mock_resp.headers = {"content-type": "application/rss+xml"}
            mock_resp.text = rss
            mock_resp.raise_for_status = lambda: None
            mock_client_cls.return_value.__enter__ = lambda s: mock_client_cls.return_value
            mock_client_cls.return_value.__exit__ = MagicMock(return_value=False)
            mock_client_cls.return_value.get = MagicMock(return_value=mock_resp)

            resp = client.post("/agent/knowledge/fetch-url", headers=headers, json={"url": "https://example.com/feed.rss"})

        assert resp.status_code == 201
        assert resp.json()["record_count"] == 2

    def test_fetch_url_http_error_returns_422(self, client, agent):
        _, _, headers = agent
        with patch("src.infrastructure.ingestion.url_fetcher.httpx.Client") as mock_client_cls:
            mock_resp = MagicMock()
            mock_resp.raise_for_status.side_effect = httpx.HTTPStatusError(
                "404", request=MagicMock(), response=MagicMock()
            )
            mock_client_cls.return_value.__enter__ = lambda s: mock_client_cls.return_value
            mock_client_cls.return_value.__exit__ = MagicMock(return_value=False)
            mock_client_cls.return_value.get = MagicMock(return_value=mock_resp)

            resp = client.post("/agent/knowledge/fetch-url", headers=headers, json={"url": "https://example.com/404"})

        assert resp.status_code == 422

    def test_fetch_url_empty_content_returns_422(self, client, agent):
        _, _, headers = agent
        with patch("src.infrastructure.ingestion.url_fetcher.httpx.Client") as mock_client_cls:
            mock_resp = MagicMock()
            mock_resp.headers = {"content-type": "text/html"}
            mock_resp.text = "<html><body></body></html>"
            mock_resp.raise_for_status = lambda: None
            mock_client_cls.return_value.__enter__ = lambda s: mock_client_cls.return_value
            mock_client_cls.return_value.__exit__ = MagicMock(return_value=False)
            mock_client_cls.return_value.get = MagicMock(return_value=mock_resp)

            resp = client.post("/agent/knowledge/fetch-url", headers=headers, json={"url": "https://example.com/empty"})

        assert resp.status_code == 422

    def test_fetch_url_filename_is_hostname(self, client, agent):
        _, _, headers = agent
        html = "<html><body><p>Conteúdo.</p></body></html>"
        with patch("src.infrastructure.ingestion.url_fetcher.httpx.Client") as mock_client_cls:
            mock_resp = MagicMock()
            mock_resp.headers = {"content-type": "text/html"}
            mock_resp.text = html
            mock_resp.raise_for_status = lambda: None
            mock_client_cls.return_value.__enter__ = lambda s: mock_client_cls.return_value
            mock_client_cls.return_value.__exit__ = MagicMock(return_value=False)
            mock_client_cls.return_value.get = MagicMock(return_value=mock_resp)

            resp = client.post("/agent/knowledge/fetch-url", headers=headers, json={"url": "https://docs.example.com/page"})

        assert resp.status_code == 201
        assert resp.json()["filename"] == "docs.example.com"

    def test_fetched_url_appears_in_knowledge_list(self, client, agent):
        _, _, headers = agent
        html = "<html><body><p>Texto indexado.</p></body></html>"
        with patch("src.infrastructure.ingestion.url_fetcher.httpx.Client") as mock_client_cls:
            mock_resp = MagicMock()
            mock_resp.headers = {"content-type": "text/html"}
            mock_resp.text = html
            mock_resp.raise_for_status = lambda: None
            mock_client_cls.return_value.__enter__ = lambda s: mock_client_cls.return_value
            mock_client_cls.return_value.__exit__ = MagicMock(return_value=False)
            mock_client_cls.return_value.get = MagicMock(return_value=mock_resp)

            client.post("/agent/knowledge/fetch-url", headers=headers, json={"url": "https://lista.example.com"})

        files = client.get("/agent/knowledge", headers=headers).json()["files"]
        types = [f["file_type"] for f in files]
        assert "url" in types
